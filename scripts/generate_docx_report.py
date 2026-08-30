#!/usr/bin/env python3
"""
Generates the NetSage AI Project Report (.docx) matching Deepak's friend format exactly.
Same 18-section structure, same table styles, same heading hierarchy.
"""

from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

ROOT = Path(__file__).parent.parent
OUTPUT_DOCX = ROOT / "NetSage_AI_Project_Report.docx"

def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def add_styled_table(doc, headers, rows, header_bg="1F2937", alt_bg="F9FAFB"):
    """Create a clean table matching Deepak's report style."""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, header_bg)
    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, text in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            p = cell.paragraphs[0]
            r = p.add_run(str(text))
            r.font.size = Pt(10)
            if ri % 2 == 0:
                set_cell_shading(cell, alt_bg)
    return table

def create_report():
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════
    # COVER / TITLE BLOCK (matching Deepak's format exactly)
    # ═══════════════════════════════════════════════════════════
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover.add_run("CISCO VIP INTERNSHIP PROJECT")
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style = doc.styles['Title']
    r = title.add_run("NetSage AI")
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("AI-Assisted Network Troubleshooting with Human Review")
    r.font.size = Pt(14)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Metadata table (matching Deepak's Table 0)
    add_styled_table(doc,
        ["Field", "Details"],
        [
            ["Technology / Domain", "Artificial Intelligence"],
            ["Project", "Applied AI + Network Troubleshooting"],
            ["Project Type", "AI-assisted troubleshooting prototype"],
            ["Team Size", "[Your Team Size] (college submission group)"],
            ["Prepared for", "Cisco VIP Internship Evaluation"],
            ["Date", "August 2026"],
        ]
    )
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════
    # 1. Executive Summary
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "NetSage AI is an AI-assisted troubleshooting helper designed for Cisco-style network lab problems. "
        "The system accepts a network symptom, topology information, and show-command evidence; performs "
        "deterministic checks for common configuration mistakes; generates a structured AI diagnosis using "
        "a Large Language Model (Groq API); and routes every diagnosis through a mandatory human review gate "
        "before any corrective action is approved."
    )
    doc.add_paragraph(
        "The project contains a 50-case troubleshooting dataset covering VLAN, default gateway, DHCP, DNS, "
        "routing, ACL, NAT, wireless, STP, EtherChannel, HSRP, IPv6, Security, and QoS scenarios. "
        "A structured prompt library defines the required AI response fields, while a Python rule checker "
        "provides deterministic evidence independently of the AI."
    )
    doc.add_paragraph(
        "The current project evaluation contains 50 AI diagnosis records. Using exact text matching against "
        "the reference fault field, 41 of 50 diagnoses matched exactly, giving an exact-match rate of 82%. "
        "6 cases required human editing and 3 were fully rejected. This metric is specific to the project's "
        "reference labels and does not represent a general model accuracy claim."
    )

    # ═══════════════════════════════════════════════════════════
    # 2. Problem Statement
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("2. Problem Statement", level=1)
    doc.add_paragraph(
        "Junior network engineers may know individual commands but can struggle to connect a symptom to the "
        "actual root cause. A single connectivity symptom may arise from VLAN, routing, DHCP, DNS, ACL, NAT, "
        "gateway, wireless, STP, HSRP, or IPv6 configuration issues."
    )
    doc.add_paragraph(
        "NetSage AI addresses this by combining network evidence, deterministic validation, AI-assisted "
        "reasoning, and mandatory human review."
    )

    # ═══════════════════════════════════════════════════════════
    # 3. Objectives
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("3. Objectives", level=1)
    objectives = [
        "Create a structured troubleshooting dataset with at least 30 cases (achieved: 50 cases).",
        "Use structured prompts to generate evidence-backed AI diagnoses.",
        "Implement deterministic checks for common network configuration mistakes.",
        "Record AI confidence, evidence, next diagnostic command, and fix steps.",
        "Require a human review decision before a proposed fix is accepted.",
        "Document AI uncertainty and correction cases through Responsible AI review.",
        "Provide a dashboard for case themes, severity, confidence, and review status.",
        "Deploy a live interactive web console accessible via public URL.",
        "Demonstrate the end-to-end troubleshooting workflow in a short video.",
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    # ═══════════════════════════════════════════════════════════
    # 4. Cisco Requirement Mapping
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("4. Cisco Requirement Mapping", level=1)
    add_styled_table(doc,
        ["Requirement", "NetSage AI Implementation", "Status"],
        [
            ["30+ troubleshooting cases", "50 cases in data/cases.csv", "Completed"],
            ["Evidence per case", "Symptom, topology, show outputs, expected fault, OSI layer, concept, severity", "Completed"],
            ["Structured AI prompts", "prompts/diagnose_prompt.md with JSON output and worked examples", "Completed"],
            ["Python rule checker", "Six deterministic validation categories in scripts/rule_checker.py", "Completed"],
            ["Dashboard", "Streamlit dashboard + Excel dashboard (outputs/dashboard.xlsx)", "Completed"],
            ["Responsible AI", "Nine cases selected for structured human validation (exceeds 5 minimum)", "Completed"],
            ["Live deployment", "Streamlit Cloud: packettraceraidiagnoser.streamlit.app", "Completed"],
            ["Demo video", "Integrated workflow demonstration — recorded walkthrough", "Completed"],
        ]
    )

    # ═══════════════════════════════════════════════════════════
    # 5. Dataset Design
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("5. Dataset Design", level=1)
    doc.add_paragraph(
        "The dataset contains 50 cases distributed across twelve troubleshooting categories. Each case "
        "includes the evidence needed for diagnosis and a reference fault used for project evaluation."
    )

    doc.add_heading("5.1 Dataset Fields", level=2)
    add_styled_table(doc,
        ["Field", "Purpose"],
        [
            ["case_id", "Unique case identifier (C001–C050)"],
            ["category", "Troubleshooting category (e.g. VLAN, Routing, DHCP)"],
            ["symptom", "Observed network problem"],
            ["topology_note", "Relevant network topology/context"],
            ["show_output", "Command output used as evidence"],
            ["expected_fault", "Reference diagnosis (ground truth)"],
            ["osi_layer", "Relevant OSI layer (Layer 1–7)"],
            ["concept_tag", "Troubleshooting concept/category"],
            ["severity", "Impact classification (Critical/High/Medium/Low)"],
        ]
    )

    doc.add_paragraph()
    doc.add_heading("5.2 Category Distribution", level=2)
    add_styled_table(doc,
        ["Category", "Number of Cases"],
        [
            ["VLAN assignment / trunk allowance", "6"],
            ["Default gateway / inter-VLAN routing", "4"],
            ["DHCP pool / relay / exhaustion", "4"],
            ["DNS server / forwarding / records", "4"],
            ["Static routing / OSPF / BGP", "6"],
            ["ACL blocking / ordering / direction", "4"],
            ["NAT inside/outside / overload / pool", "3"],
            ["STP root bridge / BPDU Guard / PortFast", "3"],
            ["EtherChannel LACP / load balance", "2"],
            ["HSRP preempt / split-brain", "2"],
            ["IPv6 unicast-routing / SLAAC", "2"],
            ["Wireless PSK / SSID VLAN / interference", "3"],
            ["Security port-security / SSH / AAA", "3"],
            ["QoS LLQ / CoS trust", "2"],
            ["Advanced routing loop / eBGP multihop", "2"],
            ["Total", "50"],
        ]
    )

    # ═══════════════════════════════════════════════════════════
    # 6. AI Prompt Design
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("6. AI Prompt Design", level=1)
    doc.add_paragraph(
        "The diagnosis prompt instructs the AI to use only supplied evidence, avoid inventing facts, "
        "state uncertainty when evidence is insufficient, and return structured JSON."
    )
    add_styled_table(doc,
        ["Required AI Field", "Purpose"],
        [
            ["root_cause", "Most likely network fault (one sentence)"],
            ["confidence", "High, Medium, or Low"],
            ["osi_layer", "Layer associated with the diagnosis"],
            ["evidence", "Specific evidence from show-command output supporting the diagnosis"],
            ["next_command", "Next useful diagnostic CLI command to confirm"],
            ["fix_steps", "Proposed remediation steps (ordered list)"],
        ]
    )

    # ═══════════════════════════════════════════════════════════
    # 7. Deterministic Rule Checker
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("7. Deterministic Rule Checker", level=1)
    doc.add_paragraph(
        "The Python checker is independent of the AI diagnosis. Its role is to produce objective "
        "configuration evidence that can support or challenge the AI."
    )
    add_styled_table(doc,
        ["Check", "Purpose"],
        [
            ["Duplicate IP", "Detect repeated IPv4 addresses in show-command output"],
            ["Subnet / network mismatch", "Check whether the host belongs to the expected network"],
            ["Gateway mismatch", "Check whether the default gateway belongs to the host network"],
            ["Interface status", "Identify down or administratively down interfaces"],
            ["VLAN existence", "Check whether the expected VLAN appears in supplied VLAN evidence"],
            ["Route existence", "Check whether a destination route appears in supplied routing evidence"],
        ]
    )

    # ═══════════════════════════════════════════════════════════
    # 8. System Architecture
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("8. System Architecture", level=1)
    doc.add_paragraph(
        "The system separates probabilistic AI reasoning from deterministic checks and human decision-making. "
        "The AI does not autonomously modify a network device."
    )
    # Architecture flow (matching Deepak's arrow diagram)
    flow = doc.add_paragraph()
    flow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = flow.add_run(
        "Troubleshooting Case → Evidence Validator → Python Rule Checker + AI Diagnosis → "
        "Human Review Gate → Approved Fix + Verification → Evaluation/Dashboard"
    )
    r.font.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

    # ═══════════════════════════════════════════════════════════
    # 9. Integrated Prototype
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("9. Integrated Prototype", level=1)
    doc.add_paragraph(
        "The integrated Streamlit prototype allows a reviewer to select a case, inspect the symptom and "
        "topology, view show-command evidence, inspect automated rule indicators, review the AI diagnosis, "
        "and record a human decision as Pending, Accepted, Edited, or Rejected. The 5-step gated pipeline "
        "enforces that Tab 05 Verification remains locked until the human review verdict is saved in Tab 04."
    )
    doc.add_paragraph(
        "The application stores review records separately in data/review_log.csv so that human decisions "
        "are not implicitly inferred from AI output."
    )
    doc.add_paragraph(
        "The live prototype is deployed at: https://packettraceraidiagnoser-pyrbpk7tdzsbppwfmadsqf.streamlit.app/"
    )

    # ═══════════════════════════════════════════════════════════
    # 10. AI Evaluation
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("10. AI Evaluation", level=1)
    doc.add_paragraph(
        "The evaluation comparison is intentionally simple and transparent. A production evaluation should "
        "also measure semantic correctness, evidence grounding, command quality, and safety of the proposed "
        "fix rather than relying only on exact text."
    )
    add_styled_table(doc,
        ["Metric", "Current Result"],
        [
            ["Total cases", "50"],
            ["Exact root-cause matches (Accepted)", "41"],
            ["Human-edited diagnoses", "6"],
            ["Human-rejected diagnoses", "3"],
            ["Exact root-cause match rate", "82%"],
            ["Evaluation type", "Exact text comparison against project reference labels"],
            ["General model accuracy claim?", "No — metric applies only to project reference labels"],
        ]
    )

    # ═══════════════════════════════════════════════════════════
    # 11. Responsible AI and Human Oversight
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("11. Responsible AI and Human Oversight", level=1)
    doc.add_paragraph(
        "Human review is a mandatory gate. Nine AI diagnosis records are included in the Human Review "
        "workflow because they contain uncertainty or require clarification. The application provides "
        "explicit reviewer fields for the decision, reviewer name, and reviewer notes. The project exceeds "
        "the minimum requirement of 5 corrected cases."
    )
    add_styled_table(doc,
        ["Case", "Reason Selected for Review"],
        [
            ["C003", "Diagnosis is directionally correct but native VLAN standard needs clarification."],
            ["C004", "Port was already in VLAN 30. Fault was a stale DHCP lease from VLAN 10."],
            ["C016", "DNS server was reachable but not authoritative for internal domain."],
            ["C018", "Route existed but next-hop 10.0.0.6 was invalid for the /30 link."],
            ["C023", "Same ACL applied inbound and outbound, hitting implicit deny. Removed outbound ACL."],
            ["C029", "SSID mapped to internal VLAN 10 instead of Guest VLAN 30 — security risk."],
            ["C033", "PortFast trunk enabled on uplink created L2 loop. Removed PortFast trunk."],
            ["C037", "HSRP hellos were blocked between peers causing split-brain active contention."],
            ["C046", "Mutual static routes between R1 and R2 created a 2-way routing loop."],
        ]
    )

    # ═══════════════════════════════════════════════════════════
    # 12. Dashboard
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("12. Dashboard", level=1)
    doc.add_paragraph(
        "The dashboard is designed to show total cases, review status, accepted/edited/rejected decisions, "
        "AI confidence distribution, issue types, severity breakdown, and AI-human agreement based on "
        "recorded human-review decisions. The web dashboard is built with Streamlit and Plotly. An Excel "
        "dashboard (outputs/dashboard.xlsx) is also generated with native pivot charts."
    )

    # ═══════════════════════════════════════════════════════════
    # 13. Demonstration Plan
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("13. Demonstration Plan", level=1)
    demo_steps = [
        "Introduce the network troubleshooting problem and NetSage AI.",
        "Show the architecture (3-tier: Rule Checker + AI Diagnosis + Human Gate).",
        "Open the case dataset and select a broken case.",
        "Display symptom, topology, and show-command evidence.",
        "Run/show deterministic rule checks.",
        "Display the structured AI diagnosis.",
        "Demonstrate human review and the Accept/Edit/Reject decision controls.",
        "Show that Tab 05 Verification unlocks only after human approval.",
        "Show the proposed fix and verification approach.",
        "Demonstrate the Live Diagnose feature with on-demand AI diagnosis.",
        "Open the dashboard and explain the evaluation metrics.",
    ]
    for step in demo_steps:
        doc.add_paragraph(step, style='List Bullet')

    # ═══════════════════════════════════════════════════════════
    # 14. Individual Contribution
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("14. Individual Contribution", level=1)
    doc.add_paragraph(
        "The team members contributed collaboratively, with distinct primary technical responsibilities "
        "documented in their individual summary reports."
    )
    add_styled_table(doc,
        ["Contribution Area", "Student's Primary Work"],
        [
            ["AI architecture & prompt engineering", "[Student Name] — Defined the AI-assisted troubleshooting workflow, structured diagnosis prompt, and Groq API integration."],
            ["Application engineering & dashboard", "[Student Name] — Contributed to the Streamlit application, interactive case workflow, 5-step gated pipeline, and Plotly dashboard."],
            ["Testing, Responsible AI & QA", "[Student Name] — Contributed to functional testing, case-level consistency checks, and responsible AI audit log documentation."],
            ["Dataset engineering & AI evaluation", "[Student Name] — Contributed to organizing and validating the 50-case dataset, evaluation comparison, and Excel report generation."],
            ["Integration, testing & review", "All members — Contributed to integration testing, problem solving, final review, and deployment."],
        ]
    )

    # ═══════════════════════════════════════════════════════════
    # 15. Limitations and Future Improvements
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("15. Limitations and Future Improvements", level=1)
    limitations = [
        "The current prototype is designed for Cisco-style lab scenarios rather than direct production network control.",
        "The rule checker intentionally covers a limited set of deterministic checks.",
        "The current exact-match metric does not fully capture semantic correctness.",
        "Human-review decisions should be documented through the Human Review workflow before Responsible AI outcomes are reported externally.",
        "A future version could connect approved diagnostics to a controlled lab verification environment.",
        "Future evaluation could include evidence-grounding and fix-safety scoring.",
    ]
    for lim in limitations:
        doc.add_paragraph(lim, style='List Bullet')

    # ═══════════════════════════════════════════════════════════
    # 16. Conclusion
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("16. Conclusion", level=1)
    doc.add_paragraph(
        "NetSage AI demonstrates how AI can assist network troubleshooting without replacing human judgment. "
        "The project combines structured network evidence, deterministic configuration checks, AI-generated "
        "diagnoses, human review, and evaluation visualization into a unified system. The complete project "
        "with 50 diagnostic scenarios is deployed live and accessible for evaluation."
    )

    # ═══════════════════════════════════════════════════════════
    # 17. Project Artifacts
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("17. Project Artifacts", level=1)
    artifacts = [
        "data/cases.csv — 50 structured troubleshooting cases",
        "prompts/diagnose_prompt.md — AI system prompt with JSON schema",
        "scripts/rule_checker.py — Deterministic configuration checker",
        "scripts/ai_diagnose.py — Groq API diagnosis pipeline",
        "data/sample_ai_responses.csv — AI diagnosis results for all 50 cases",
        "data/review_log.csv — Human review audit log",
        "scripts/build_dashboard.py — Excel dashboard generator",
        "outputs/dashboard.xlsx — Generated Excel dashboard with charts",
        "docs/responsible_ai_log.md — Responsible AI correction documentation",
        "app.py — Streamlit web console (deployed on Streamlit Cloud)",
        "packet_tracer_lab/ — Packet Tracer configs (SW1, R1 broken + working)",
    ]
    for art in artifacts:
        doc.add_paragraph(art, style='List Bullet')

    # ═══════════════════════════════════════════════════════════
    # 18. Group Members and Individual Submission
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("18. Group Members and Individual Submission", level=1)
    doc.add_paragraph(
        "The shared technical project is developed for the following group:"
    )
    add_styled_table(doc,
        ["Member", "Name", "Section", "Registration ID"],
        [
            ["1", "[Your Name]", "[Your Section]", "[Your Registration ID]"],
            ["2", "[Partner 2 Name]", "[Section]", "[Registration ID]"],
            ["3", "[Partner 3 Name]", "[Section]", "[Registration ID]"],
            ["4", "[Partner 4 Name]", "[Section]", "[Registration ID]"],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "The shared project artifacts may be common to the group. Each student should use their own "
        "individual report version and describe only the work they actually performed. Do not copy "
        "another member's individual contribution section."
    )

    # Save
    doc.save(OUTPUT_DOCX)
    print(f"Report generated: {OUTPUT_DOCX}")

if __name__ == "__main__":
    create_report()
